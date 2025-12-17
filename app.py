import streamlit as st
import pandas as pd
import pdfplumber
import docx
import os
import json
from io import BytesIO
from dotenv import load_dotenv
from openai import OpenAI

# =========================
# CONFIG
# =========================
st.set_page_config(
    page_title="Trợ lý Ma Trận Đặc Tả",
    layout="wide"
)

load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

MODEL_NAME = "gpt-4.1"

# =========================
# HELPER: READ FILE
# =========================
def read_pdf(file):
    text = ""
    with pdfplumber.open(file) as pdf:
        for i, page in enumerate(pdf.pages):
            text += f"\n--- Page {i+1} ---\n"
            text += page.extract_text() or ""
    return text

def read_docx(file):
    doc = docx.Document(file)
    return "\n".join(p.text for p in doc.paragraphs)

def read_excel(file):
    df = pd.read_excel(file)
    return df.to_csv(index=False)

from io import BytesIO

def extract_text(uploaded_file):
    # Đọc file thành bytes
    file_bytes = uploaded_file.read()
    uploaded_file.seek(0)  # reset con trỏ file

    file_name = uploaded_file.name.lower()

    if file_name.endswith(".pdf"):
        text = ""
        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Trang {i+1} ---\n{page_text}"
        return text.strip()

    elif file_name.endswith(".docx"):
        doc = docx.Document(BytesIO(file_bytes))
        texts = []

        for p in doc.paragraphs:
            if p.text.strip():
                texts.append(p.text)

        # Đọc cả bảng trong Word
        for table in doc.tables:
            for row in table.rows:
                texts.append(" | ".join(cell.text for cell in row.cells))

        return "\n".join(texts).strip()

    elif file_name.endswith(".xlsx") or file_name.endswith(".xls"):
        df = pd.read_excel(BytesIO(file_bytes))
        return df.to_csv(index=False)

    else:
        return file_bytes.decode("utf-8", errors="ignore").strip()


# =========================
# UI – HEADER
# =========================
st.markdown("""
# 🧠 **TRỢ LÍ MA TRẬN ĐẶC TẢ**
_Hỗ trợ xây dựng bảng đặc tả đề kiểm tra – chuẩn khảo thí_
""")

# =========================
# SECTION 1 – DATA
# =========================
st.header("① Dữ liệu tham chiếu")

ref_files = st.file_uploader(
    "Upload tài liệu (PDF / Word / Excel / Text)",
    type=["pdf", "docx", "xlsx", "xls", "txt", "csv"],
    accept_multiple_files=True
)

ref_text = st.text_area(
    "Hoặc dán nội dung trực tiếp",
    height=200
)

text = extract_text(f)

if not text or len(text) < 50:
    st.warning(f"⚠️ File {f.name} không trích xuất được nội dung (PDF scan hoặc file rỗng)")
else:
    reference_contents.append(
        f"\n=== FILE: {f.name} ===\n{text}"
    )

    # DEBUG – xem trước 500 ký tự
    with st.expander(f"📄 Xem trước nội dung {f.name}"):
        st.text(text[:500])

# =========================
# SECTION 2 – TEMPLATE
# =========================
st.header("② Khung ma trận mẫu")

template_file = st.file_uploader(
    "Upload file mẫu",
    type=["pdf", "docx", "xlsx", "xls", "txt", "csv"],
    accept_multiple_files=False
)

default_template = (
    "STT, Nội dung kiến thức, Đơn vị kiến thức, "
    "Chuẩn cần đánh giá, Nhận biết, Thông hiểu, "
    "Vận dụng, Vận dụng cao, Tổng số câu, Ghi chú"
)

template_text = st.text_area(
    "Khung cột ma trận",
    value=default_template,
    height=150
)

if template_file:
    with st.spinner("Đang đọc file mẫu..."):
        template_text += "\n\n" + extract_text(template_file)

# =========================
# SECTION 3 – GENERATE
# =========================
st.header("③ Tạo ma trận bằng AI")

if st.button("🚀 TẠO MA TRẬN ĐẶC TẢ", use_container_width=True):

    if not reference_contents and not ref_text.strip():
        st.error("❌ Chưa có dữ liệu tham chiếu")
        st.stop()

    with st.spinner("GPT-4.1 đang phân tích và xây dựng ma trận..."):

        system_prompt = """
Bạn là CHUYÊN GIA KHẢO THÍ.

NHIỆM VỤ:
- Phân tích dữ liệu môn học
- Tạo BẢNG MA TRẬN ĐẶC TẢ

⚠️ QUY TẮC BẮT BUỘC:
1. Chỉ trả về JSON
2. Không markdown
3. Không giải thích
4. TẤT CẢ giá trị trong rows PHẢI LÀ STRING
5. Không number, không null
"""

        user_prompt = f"""
=== KHUNG MA TRẬN ===
{template_text}

=== DỮ LIỆU THAM CHIẾU ===
{ref_text}

{"".join(reference_contents)}
"""

        response = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            response_format={
                "type": "json_schema",
                "json_schema": {
                    "name": "matrix_spec",
                    "schema": {
                        "type": "object",
                        "properties": {
                            "headers": {
                                "type": "array",
                                "items": {"type": "string"}
                            },
                            "rows": {
                                "type": "array",
                                "items": {
                                    "type": "array",
                                    "items": {"type": "string"}
                                }
                            }
                        },
                        "required": ["headers", "rows"]
                    }
                }
            },
            temperature=0.2
        )

        try:
            result = json.loads(response.choices[0].message.content)
            df = pd.DataFrame(result["rows"], columns=result["headers"])

            st.success("✅ Tạo ma trận thành công")
            st.dataframe(df, use_container_width=True)

            csv = df.to_csv(index=False).encode("utf-8-sig")
            st.download_button(
                "⬇️ Tải file CSV",
                csv,
                "Ma_Tran_Dac_Ta.csv",
                "text/csv"
            )

        except Exception as e:
            st.error("❌ GPT-4.1 trả dữ liệu lỗi")
            st.code(response.choices[0].message.content)
